"""Docling adapter. Parser details do not leak into upload or retrieval code."""

import hashlib
import re
from dataclasses import dataclass, field
from pathlib import Path


class DocumentParseError(ValueError):
    pass


@dataclass(frozen=True)
class ParsedDocument:
    markdown: str
    content_hash: str
    heading2_pages: tuple[tuple[str, int], ...] = field(default_factory=tuple)
    total_pages: int | None = None


def _docx_heading_structure(
    path: Path,
) -> tuple[tuple[tuple[int, str, int], ...], int]:
    """Read ordered Word headings and their explicit DOCX pages."""
    from docx import Document

    document = Document(path)
    page = 1
    headings: list[tuple[int, str, int]] = []
    for block in document.iter_inner_content():
        paragraphs = (
            [block]
            if hasattr(block, "_p")
            else [
                paragraph
                for row in block.rows
                for cell in row.cells
                for paragraph in cell.paragraphs
            ]
        )
        for paragraph in paragraphs:
            style_name = (paragraph.style.name or "").casefold().replace(" ", "")
            style_id = (paragraph.style.style_id or "").casefold().replace(" ", "")
            text = paragraph.text.strip()
            style_match = re.fullmatch(r"heading([1-6])", style_name)
            style_match = style_match or re.fullmatch(r"heading([1-6])", style_id)
            if text and style_match:
                headings.append((int(style_match.group(1)), text, page))
            page_breaks = len(paragraph._p.xpath('.//w:br[@w:type="page"]'))
            rendered_breaks = len(paragraph._p.xpath(".//w:lastRenderedPageBreak"))
            page += max(page_breaks, rendered_breaks)
    return tuple(headings), page


def _markdown_heading_text(line: str) -> str:
    candidate = re.sub(r"^#{1,6}\s+", "", line.strip())
    candidate = re.sub(r"^(?:\*\*|__)(.*)(?:\*\*|__)$", r"\1", candidate)
    return candidate.strip()


def _normalize_docx_heading2(markdown: str, headings: tuple[tuple[int, str, int], ...]) -> str:
    """Preserve Word heading levels so Heading 2 remains the boundary marker."""
    if not headings:
        return markdown
    lines = markdown.splitlines()
    cursor = 0
    for level, title, _page in headings:
        for index in range(cursor, len(lines)):
            if _markdown_heading_text(lines[index]) == title:
                lines[index] = f"{'#' * level} {title}"
                cursor = index + 1
                break
        else:
            raise DocumentParseError(f"DOCX heading could not be preserved in Markdown: {title}")
    return "\n".join(lines)


def parse_to_markdown(path: Path, filename: str) -> ParsedDocument:
    extension = Path(filename).suffix.lower()
    heading2_pages: tuple[tuple[str, int], ...] = ()
    total_pages = None
    try:
        if extension in {".md", ".txt"}:
            markdown = path.read_text(encoding="utf-8")
        else:
            from docling.document_converter import DocumentConverter

            result = DocumentConverter().convert(path)
            markdown = result.document.export_to_markdown()
            if extension == ".docx":
                headings, total_pages = _docx_heading_structure(path)
                heading2_pages = tuple(
                    (title, page) for level, title, page in headings if level == 2
                )
                markdown = _normalize_docx_heading2(markdown, headings)
    except UnicodeDecodeError as exc:
        raise DocumentParseError("text document is not valid UTF-8") from exc
    except Exception as exc:
        raise DocumentParseError("document parsing failed") from exc
    markdown = markdown.replace("\r\n", "\n").replace("\r", "\n")
    if not markdown.strip():
        raise DocumentParseError("document did not contain extractable text")
    return ParsedDocument(
        markdown=markdown,
        content_hash=hashlib.sha256(markdown.encode()).hexdigest(),
        heading2_pages=heading2_pages,
        total_pages=total_pages,
    )
