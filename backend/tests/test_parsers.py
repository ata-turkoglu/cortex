import pytest

from app.ingestion.parsers import DocumentParseError, parse_to_markdown


def test_text_document_is_normalized_to_markdown(tmp_path):
    source = tmp_path / "notes.txt"
    source.write_bytes(b"First\r\nSecond\rThird")
    parsed = parse_to_markdown(source, source.name)
    assert parsed.markdown == "First\nSecond\nThird"
    assert len(parsed.content_hash) == 64


def test_invalid_text_is_a_structured_parse_failure(tmp_path):
    source = tmp_path / "bad.txt"
    source.write_bytes(b"\xff\xfe")
    with pytest.raises(DocumentParseError, match="UTF-8"):
        parse_to_markdown(source, source.name)


def test_docling_converts_docx_to_markdown(tmp_path):
    from docx import Document

    source = tmp_path / "notes.docx"
    document = Document()
    document.add_heading("Başlık", level=1)
    document.add_paragraph("İçerik")
    document.save(source)
    parsed = parse_to_markdown(source, source.name)
    assert "Başlık" in parsed.markdown
    assert "İçerik" in parsed.markdown


def test_docx_heading2_is_preserved_and_mapped_to_explicit_pages(tmp_path):
    from docx import Document

    source = tmp_path / "MERTER B.docx"
    document = Document()
    document.add_heading("MERTER B", level=1)
    document.add_heading("B-1/a", level=2)
    document.add_paragraph("Birinci belge")
    document.add_page_break()
    document.add_heading("B-2/i", level=2)
    document.add_paragraph("İkinci belge")
    document.save(source)
    parsed = parse_to_markdown(source, source.name)
    assert "# MERTER B" in parsed.markdown
    assert "## B-1/a" in parsed.markdown
    assert "## B-2/i" in parsed.markdown
    assert parsed.heading2_pages == (("B-1/a", 1), ("B-2/i", 2))
    assert parsed.total_pages == 2
