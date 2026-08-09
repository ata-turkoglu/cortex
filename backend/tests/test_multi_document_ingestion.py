import io
import json

from fastapi.testclient import TestClient
from sqlalchemy import create_engine, event, select
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

from app.api.workspaces import get_session
from app.core.config import Settings
from app.main import app
from app.models import Base, Chunk, LogicalDocument


def test_logical_document_splitter_uses_only_level_two_headings():
    from app.ingestion.logical_documents import detect_logical_documents
    from app.ingestion.parsers import ParsedDocument

    markdown = """# Archive container

B-9/x
This plain archive-like code is part of the preamble.

## S.17
First logical document.
### B-10/z
This level-three heading stays in S.17.

## Future Archive Ω 2026
Second logical document.
B-11/a
This plain code is content, not a boundary.
"""
    parsed = ParsedDocument(markdown=markdown, content_hash="hash")

    drafts = detect_logical_documents(parsed, "archive.docx", "document")

    assert [draft.document_code for draft in drafts] == [
        "S.17",
        "Future Archive Ω 2026",
    ]
    assert drafts[0].markdown.startswith("## S.17")
    assert "### B-10/z" in drafts[0].markdown
    assert "B-11/a" in drafts[1].markdown


def test_archive_like_plain_text_does_not_create_logical_boundaries():
    from app.ingestion.logical_documents import detect_logical_documents
    from app.ingestion.parsers import ParsedDocument

    markdown = "# MERTER B\n\nB-1/a\nFirst\n\nA-21/b\nSecond"
    drafts = detect_logical_documents(
        ParsedDocument(markdown=markdown, content_hash="hash"),
        "MERTER B.docx",
        "document",
    )

    assert len(drafts) == 1
    assert drafts[0].document_code == "MERTER B"


def _foreign_keys(connection, _):
    connection.execute("PRAGMA foreign_keys=ON")


def _multi_document_docx() -> bytes:
    from docx import Document

    output = io.BytesIO()
    document = Document()
    document.add_heading("MERTER B", level=1)
    document.add_heading("B-1/a", level=2)
    document.add_paragraph("Ahmet ilk tapu belgesinde mal sahibidir.")
    document.add_page_break()
    document.add_heading("B-2/i", level=2)
    document.add_paragraph("Hüseyin Hüsnü Subaşı ikinci tapu belgesinde mal sahibidir.")
    document.save(output)
    return output.getvalue()


def test_docx_upload_preserves_logical_documents_chunks_and_diagnostics(tmp_path, monkeypatch):
    engine = create_engine(
        "sqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    event.listen(engine, "connect", _foreign_keys)
    Base.metadata.create_all(engine)
    session_local = sessionmaker(bind=engine)
    settings = Settings(data_path=tmp_path)
    monkeypatch.setattr("app.api.uploads.get_settings", lambda: settings)
    monkeypatch.setattr("app.core.workspaces.get_settings", lambda: settings)

    def session_override():
        session = session_local()
        try:
            yield session
            session.commit()
        finally:
            session.close()

    app.dependency_overrides[get_session] = session_override
    try:
        client = TestClient(app)
        workspace = client.post(
            "/api/v1/workspaces", json={"name": "Logical", "slug": "logical"}
        ).json()
        response = client.post(
            f"/api/v1/workspaces/{workspace['id']}/uploads",
            files={
                "files": (
                    "MERTER B.docx",
                    _multi_document_docx(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert response.status_code == 201, response.text
        uploaded = response.json()["uploads"][0]
        assert uploaded["logical_document_count"] == 2
        assert [item["document_code"] for item in uploaded["logical_documents"]] == [
            "B-1/a",
            "B-2/i",
        ]

        session = session_local()
        logical = session.scalars(select(LogicalDocument).order_by(LogicalDocument.ordinal)).all()
        chunks = session.scalars(select(Chunk).order_by(Chunk.ordinal)).all()
        assert len(logical) == 2
        assert {chunk.logical_document_id for chunk in chunks} == {item.id for item in logical}
        for chunk in chunks:
            metadata = json.loads(chunk.metadata_json)
            owner = next(item for item in logical if item.id == chunk.logical_document_id)
            assert metadata["document_id"] == owner.id
            assert metadata["document_code"] == owner.document_code
            assert metadata["source_original"] == "MERTER B.docx"
        session.close()

        diagnostic = client.get(
            f"/api/v1/workspaces/{workspace['id']}/ingestion-diagnostics/{uploaded['document_id']}",
            params={"query": "Hüseyin Hüsnü Subaşı"},
        )
        assert diagnostic.status_code == 200, diagnostic.text
        body = diagnostic.json()
        assert [item["document_code"] for item in body["logical_documents"]] == [
            "B-1/a",
            "B-2/i",
        ]
        assert body["logical_documents"][1]["page_start"] == 2
        assert body["logical_documents"][1]["chunks"]
        assert [item["document_code"] for item in body["retrieved_documents"]] == ["B-2/i"]
    finally:
        app.dependency_overrides.clear()


def test_graphrag_materializes_one_input_per_logical_document(tmp_path):
    from datetime import UTC, datetime
    from uuid import uuid4

    from app.graphrag.input import GraphRAGInputMaterializer
    from app.models import Document, DocumentVersion, Workspace

    session = sessionmaker(bind=create_engine("sqlite:///:memory:"))()
    Base.metadata.create_all(session.bind)
    now, workspace_id = datetime.now(UTC), str(uuid4())
    document_id, version_id = str(uuid4()), str(uuid4())
    session.add(
        Workspace(
            id=workspace_id,
            slug="graph-logical",
            name="Graph logical",
            state="active",
            created_at=now,
            updated_at=now,
        )
    )
    session.add(
        Document(
            id=document_id,
            workspace_id=workspace_id,
            title="MERTER B.docx",
            active_version_id=version_id,
            created_at=now,
            updated_at=now,
        )
    )
    session.add(
        DocumentVersion(
            id=version_id,
            workspace_id=workspace_id,
            document_id=document_id,
            version_number=1,
            source_hash="hash",
            source_path="source.docx",
            source_filename="MERTER B.docx",
            size_bytes=1,
            state="ready",
            created_at=now,
        )
    )
    logical_ids = []
    for ordinal, code in enumerate(("B-1/a", "B-2/i")):
        logical_id = str(uuid4())
        logical_ids.append(logical_id)
        session.add(
            LogicalDocument(
                id=logical_id,
                workspace_id=workspace_id,
                source_document_id=document_id,
                document_version_id=version_id,
                ordinal=ordinal,
                document_code=code,
                title=code,
                document_type="tapu",
                source_original="MERTER B.docx",
                page_start=ordinal + 1,
                page_end=ordinal + 1,
                normalized_content=f"## {code}\n\nİçerik",
                created_at=now,
            )
        )
    session.commit()
    materializer = GraphRAGInputMaterializer(tmp_path)
    stale = tmp_path / "graph" / "input" / "source-level-old.md"
    stale.parent.mkdir(parents=True)
    stale.write_text("all source chunks", encoding="utf-8")
    manifest = materializer.materialize(session, workspace_id, tmp_path / "graph")
    assert set(manifest.logical_document_ids) == set(logical_ids)
    assert len(list(manifest.input_root.glob("*.md"))) == 2
    assert not stale.exists()
    second = (manifest.input_root / f"{logical_ids[1]}.md").read_text(encoding="utf-8")
    assert "# B-2/i" in second
    assert "Source original: MERTER B.docx" in second


def test_graphrag_entity_provenance_resolves_through_text_units(tmp_path):
    from app.graphrag.adapter import GraphRAGAdapter

    logical_id = "11111111-1111-1111-1111-111111111111"
    (tmp_path / "documents.json").write_text(
        json.dumps(
            [
                {
                    "id": "graph-document",
                    "text": f"Logical document ID: {logical_id}\nDocument code: B-2/i",
                    "attributes": {},
                }
            ]
        ),
        encoding="utf-8",
    )
    (tmp_path / "text_units.json").write_text(
        json.dumps(
            [
                {
                    "id": "text-unit",
                    "text": "Hüseyin Hüsnü Subaşı",
                    "attributes": {"document_ids": '["graph-document"]'},
                }
            ]
        ),
        encoding="utf-8",
    )
    (tmp_path / "entities.json").write_text(
        json.dumps(
            [
                {
                    "id": "entity",
                    "text": "Hüseyin Hüsnü Subaşı",
                    "attributes": {"text_unit_ids": '["text-unit"]'},
                }
            ]
        ),
        encoding="utf-8",
    )
    adapter = GraphRAGAdapter("workspace", tmp_path)
    entity = adapter.load_artifacts("entities")[0]
    assert adapter.logical_document_ids_for(entity) == (logical_id,)
